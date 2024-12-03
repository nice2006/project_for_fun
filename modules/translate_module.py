#! python3
# translate_module.py - 翻译模块
# -*- coding: utf-8 -*-
# time: 2024/11/24

from googletrans import Translator
from docx import Document
import os

def translate_text(text, target_language):
    """
    翻译输入文本。
    :param text: 需要翻译的文本。
    :param target_language: 翻译目标语言代码。
    :return: 翻译结果。
    """
    translator = Translator()
    detected_lang = translator.detect(text).lang  # 自动检测语言
    translation = translator.translate(text, src=detected_lang, dest=target_language)
    return translation.text, detected_lang

def translate_file(input_file, output_file, target_language):
    """
    翻译文件内容。
    :param input_file: 输入文件路径。
    :param output_file: 输出文件路径。
    :param target_language: 翻译目标语言代码。
    """
    translator = Translator()
    input_ext = os.path.splitext(input_file)[1]  # 获取文件扩展名

    if input_ext == ".docx":
        # 翻译 Word 文件
        doc = Document(input_file)
        translated_doc = Document()
        for para in doc.paragraphs:
            if para.text.strip():  # 忽略空行
                translated_para = translator.translate(para.text, dest=target_language).text
                translated_doc.add_paragraph(translated_para)
        translated_doc.save(output_file)
    elif input_ext == ".txt":
        # 翻译文本文件
        with open(input_file, 'r', encoding='utf-8') as f:
            content = f.read()
        translation = translator.translate(content, dest=target_language).text
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(translation)
    else:
        raise ValueError("不支持的文件类型，请使用 .txt 或 .docx 文件。")
